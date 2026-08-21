"""
合同管理系统 - 创建示例Word合同文档
这些文档作为测试本体论系统的测试案例
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

DOCUMENTS_DIR = Path(__file__).parent.parent / "documents"
DOCUMENTS_DIR.mkdir(exist_ok=True)


def create_contract_001():
    """合同001: 正常年度技术服务合同"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('技术服务合同', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 合同编号
    doc.add_paragraph('合同编号: HT-2024-001')
    doc.add_paragraph('签订日期: 2024年1月15日')
    
    # 甲乙方
    doc.add_heading('甲方（委托方）', level=1)
    doc.add_paragraph('北京科技有限公司')
    doc.add_paragraph('统一社会信用代码: 91110105MA01XXXXX')
    doc.add_paragraph('地址: 北京市朝阳区建国路88号')
    doc.add_paragraph('联系人: 张三')
    doc.add_paragraph('电话: 13800138001')
    
    doc.add_heading('乙方（服务方）', level=1)
    doc.add_paragraph('本体改造研发平台')
    doc.add_paragraph('地址: 北京市海淀区中关村')
    
    # 合同正文
    doc.add_heading('第一条 服务内容', level=2)
    doc.add_paragraph('乙方为甲方提供年度技术支持服务，包括：')
    doc.add_paragraph('1. 系统维护与监控', style='List Number')
    doc.add_paragraph('2. 技术咨询与支持', style='List Number')
    doc.add_paragraph('3. 故障排除与修复', style='List Number')
    doc.add_paragraph('4. 系统升级与优化', style='List Number')
    
    doc.add_heading('第二条 合同金额', level=2)
    doc.add_paragraph('合同总金额: 人民币500,000.00元（大写：伍拾万元整）')
    doc.add_paragraph('付款方式: 分期付款')
    doc.add_paragraph('第一期: 150,000元，合同签订后5个工作日内支付')
    doc.add_paragraph('第二期: 150,000元，2024年5月15日前支付')
    doc.add_paragraph('第三期: 100,000元，2024年8月15日前支付')
    doc.add_paragraph('第四期: 100,000元，2024年11月15日前支付')
    
    doc.add_heading('第三条 合同期限', level=2)
    doc.add_paragraph('合同有效期: 2024年1月15日至2024年12月31日')
    
    doc.add_heading('第四条 违约责任', level=2)
    doc.add_paragraph('任何一方违反本合同约定，应向对方支付合同总金额10%的违约金。')
    
    doc.add_heading('第五条 争议解决', level=2)
    doc.add_paragraph('因本合同引起的或与本合同有关的任何争议，双方应友好协商解决。协商不成的，任何一方均可向甲方所在地人民法院提起诉讼。')
    
    # 签章
    doc.add_paragraph('')
    doc.add_paragraph('甲方（盖章）: ________________')
    doc.add_paragraph('乙方（盖章）: ________________')
    doc.add_paragraph('签订日期: 2024年1月15日')
    
    doc.save(DOCUMENTS_DIR / "HT-2024-001_年度技术服务合同.docx")
    print("已创建: HT-2024-001_年度技术服务合同.docx")


def create_contract_002():
    """合同002: 软件采购合同"""
    doc = Document()
    
    title = doc.add_heading('软件采购合同', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('合同编号: HT-2024-002')
    doc.add_paragraph('签订日期: 2024年2月20日')
    
    doc.add_heading('甲方（采购方）', level=1)
    doc.add_paragraph('上海贸易有限公司')
    doc.add_paragraph('统一社会信用代码: 91310115MA1HXXXXX')
    doc.add_paragraph('地址: 上海市浦东新区陆家嘴金融中心')
    doc.add_paragraph('联系人: 李四')
    doc.add_paragraph('电话: 13800138002')
    
    doc.add_heading('乙方（供应方）', level=1)
    doc.add_paragraph('软件科技股份有限公司')
    
    doc.add_heading('第一条 采购内容', level=2)
    doc.add_paragraph('甲方向乙方采购企业级软件授权，具体如下：')
    doc.add_paragraph('产品名称: Enterprise Suite Pro')
    doc.add_paragraph('授权数量: 100用户')
    doc.add_paragraph('授权期限: 2年')
    
    doc.add_heading('第二条 合同金额', level=2)
    doc.add_paragraph('合同总金额: 人民币1,200,000.00元（大写：壹佰贰拾万元整）')
    doc.add_paragraph('付款方式: 分两期支付')
    doc.add_paragraph('第一期: 600,000元，合同签订后支付')
    doc.add_paragraph('第二期: 600,000元，6个月后支付')
    
    doc.add_heading('第三条 质量保证', level=2)
    doc.add_paragraph('乙方保证所提供的软件为正版授权，且在授权期限内提供免费升级和技术支持。')
    
    doc.add_heading('第四条 违约责任', level=2)
    doc.add_paragraph('如乙方提供的软件存在质量问题，乙方应在7个工作日内修复或更换。如无法修复，乙方应退还相应款项。')
    
    doc.add_paragraph('')
    doc.add_paragraph('甲方（盖章）: ________________')
    doc.add_paragraph('乙方（盖章）: ________________')
    doc.add_paragraph('签订日期: 2024年2月20日')
    
    doc.save(DOCUMENTS_DIR / "HT-2024-002_软件采购合同.docx")
    print("已创建: HT-2024-002_软件采购合同.docx")


def create_contract_006_risk():
    """合同006: 风险合同 - 黑名单客户"""
    doc = Document()
    
    title = doc.add_heading('销售合同', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('合同编号: HT-2024-006')
    doc.add_paragraph('签订日期: 2024年6月1日')
    
    doc.add_heading('甲方（销售方）', level=1)
    doc.add_paragraph('本体改造研发平台')
    
    doc.add_heading('乙方（购买方）', level=1)
    doc.add_paragraph('黑名单企业有限公司')
    doc.add_paragraph('统一社会信用代码: 91000000MA00XXXXX')
    doc.add_paragraph('地址: 某地某路某号')
    doc.add_paragraph('联系人: 孙八')
    doc.add_paragraph('电话: 13800138006')
    
    # 风险提示
    doc.add_heading('⚠️ 风险提示', level=2)
    doc.add_paragraph('注意：此合同涉及黑名单客户，存在高信用风险！')
    doc.add_paragraph('风险等级: 高')
    doc.add_paragraph('建议: 需要额外的信用审查和担保措施')
    
    doc.add_heading('第一条 采购内容', level=2)
    doc.add_paragraph('乙方向甲方采购产品，具体如下：')
    doc.add_paragraph('产品名称: 标准产品A')
    doc.add_paragraph('数量: 1000件')
    doc.add_paragraph('单价: 200元/件')
    
    doc.add_heading('第二条 合同金额', level=2)
    doc.add_paragraph('合同总金额: 人民币200,000.00元（大写：贰拾万元整）')
    doc.add_paragraph('付款方式: 货到付款')
    
    doc.add_heading('第三条 风险条款', level=2)
    doc.add_paragraph('鉴于乙方信用状况，甲方要求：')
    doc.add_paragraph('1. 乙方需提供银行保函')
    doc.add_paragraph('2. 付款期限缩短为15天')
    doc.add_paragraph('3. 甲方有权随时终止合同')
    
    doc.add_heading('第四条 违约责任', level=2)
    doc.add_paragraph('如乙方逾期付款，每日按未付款金额的0.5%支付违约金。')
    
    doc.add_paragraph('')
    doc.add_paragraph('甲方（盖章）: ________________')
    doc.add_paragraph('乙方（盖章）: ________________')
    doc.add_paragraph('签订日期: 2024年6月1日')
    
    doc.save(DOCUMENTS_DIR / "HT-2024-006_风险合同_黑名单客户.docx")
    print("已创建: HT-2024-006_风险合同_黑名单客户.docx")


def create_contract_007_overdue():
    """合同007: 逾期合同示例"""
    doc = Document()
    
    title = doc.add_heading('销售合同', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('合同编号: HT-2024-007')
    doc.add_paragraph('签订日期: 2024年1月20日')
    
    doc.add_heading('甲方（销售方）', level=1)
    doc.add_paragraph('本体改造研发平台')
    
    doc.add_heading('乙方（购买方）', level=1)
    doc.add_paragraph('北京科技有限公司')
    doc.add_paragraph('联系人: 张三')
    doc.add_paragraph('电话: 13800138001')
    
    doc.add_heading('第一条 采购内容', level=2)
    doc.add_paragraph('乙方向甲方采购产品，具体如下：')
    doc.add_paragraph('产品名称: 高级产品B')
    doc.add_paragraph('数量: 500件')
    doc.add_paragraph('单价: 900元/件')
    
    doc.add_heading('第二条 合同金额', level=2)
    doc.add_paragraph('合同总金额: 人民币450,000.00元（大写：肆拾伍万元整）')
    doc.add_paragraph('付款方式: 分三期支付')
    doc.add_paragraph('第一期: 150,000元，2024年3月1日前支付 ✅ 已支付')
    doc.add_paragraph('第二期: 150,000元，2024年6月1日前支付 ⚠️ 已逾期')
    doc.add_paragraph('第三期: 150,000元，2024年9月1日前支付')
    
    # 逾期风险提示
    doc.add_heading('⚠️ 逾期风险', level=2)
    doc.add_paragraph('当前状态: 第二期款项已逾期！')
    doc.add_paragraph('逾期天数: 35天')
    doc.add_paragraph('逾期金额: 150,000元')
    doc.add_paragraph('风险等级: 中')
    
    doc.add_heading('第三条 逾期处理', level=2)
    doc.add_paragraph('如乙方逾期付款，甲方将：')
    doc.add_paragraph('1. 发送催款通知')
    doc.add_paragraph('2. 暂停后续服务')
    doc.add_paragraph('3. 按合同约定收取违约金')
    
    doc.add_paragraph('')
    doc.add_paragraph('甲方（盖章）: ________________')
    doc.add_paragraph('乙方（盖章）: ________________')
    doc.add_paragraph('签订日期: 2024年1月20日')
    
    doc.save(DOCUMENTS_DIR / "HT-2024-007_逾期合同示例.docx")
    print("已创建: HT-2024-007_逾期合同示例.docx")


def create_contract_008_framework():
    """合同008: 框架合同"""
    doc = Document()
    
    title = doc.add_heading('框架合作协议', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('合同编号: HT-2024-008')
    doc.add_paragraph('签订日期: 2024年7月1日')
    
    doc.add_heading('甲方', level=1)
    doc.add_paragraph('上海贸易有限公司')
    doc.add_paragraph('统一社会信用代码: 91310115MA1HXXXXX')
    doc.add_paragraph('地址: 上海市浦东新区陆家嘴金融中心')
    doc.add_paragraph('联系人: 李四')
    
    doc.add_heading('乙方', level=1)
    doc.add_paragraph('本体改造研发平台')
    
    doc.add_heading('第一条 合作范围', level=2)
    doc.add_paragraph('甲乙双方建立长期战略合作关系，在以下领域开展合作：')
    doc.add_paragraph('1. 软件产品销售', style='List Number')
    doc.add_paragraph('2. 技术服务提供', style='List Number')
    doc.add_paragraph('3. 联合市场推广', style='List Number')
    
    doc.add_heading('第二条 合作期限', level=2)
    doc.add_paragraph('本协议有效期: 2024年7月1日至2026年6月30日')
    doc.add_paragraph('协议期满前30天，双方可协商续约。')
    
    doc.add_heading('第三条 合作金额', level=2)
    doc.add_paragraph('本协议为框架协议，具体项目金额以单独签订的子合同为准。')
    doc.add_paragraph('预计合作总金额: 人民币2,000,000.00元')
    
    doc.add_heading('第四条 权利义务', level=2)
    doc.add_paragraph('甲方权利义务：')
    doc.add_paragraph('1. 优先获得乙方最新产品信息')
    doc.add_paragraph('2. 享受协议约定的优惠政策')
    doc.add_paragraph('3. 按时支付子合同款项')
    
    doc.add_paragraph('乙方权利义务：')
    doc.add_paragraph('1. 为甲方提供专属客户经理')
    doc.add_paragraph('2. 优先响应甲方需求')
    doc.add_paragraph('3. 提供技术支持和培训')
    
    doc.add_heading('第五条 保密条款', level=2)
    doc.add_paragraph('双方对合作过程中知悉的对方商业秘密负有保密义务，未经对方书面同意，不得向第三方披露。')
    
    doc.add_paragraph('')
    doc.add_paragraph('甲方（盖章）: ________________')
    doc.add_paragraph('乙方（盖章）: ________________')
    doc.add_paragraph('签订日期: 2024年7月1日')
    
    doc.save(DOCUMENTS_DIR / "HT-2024-008_框架合作协议.docx")
    print("已创建: HT-2024-008_框架合作协议.docx")


def create_all_contracts():
    """创建所有示例合同"""
    print("=" * 50)
    print("创建示例Word合同文档")
    print("=" * 50)
    
    create_contract_001()
    create_contract_002()
    create_contract_006_risk()
    create_contract_007_overdue()
    create_contract_008_framework()
    
    print("=" * 50)
    print("所有合同文档创建完成！")
    print(f"文档目录: {DOCUMENTS_DIR}")
    print("=" * 50)


if __name__ == "__main__":
    create_all_contracts()
