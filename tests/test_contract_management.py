from __future__ import annotations

from io import BytesIO
from pathlib import Path
import sys

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from ontology_platform.contract_management import add_contract_version, compare_contract_engines, create_contract, get_contract_document, list_contracts
from ontology_platform.database import initialize_platform_db


def _contract_docx(contract_no: str = "HT-2026-100") -> bytes:
    document = Document()
    document.add_heading("软件采购合同", 0)
    document.add_paragraph(f"合同编号：{contract_no}")
    document.add_paragraph("甲方：星河科技有限公司")
    document.add_paragraph("乙方：未来软件有限公司")
    document.add_paragraph("合同金额：100000元")
    document.add_paragraph("第一条 付款安排")
    document.add_paragraph("签约后付款100000元。")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_word_document_is_required_and_version_is_downloadable(tmp_path: Path) -> None:
    database = tmp_path / "contracts.sqlite3"
    initialize_platform_db(database)

    contract = create_contract(database, "采购合同.docx", _contract_docx())

    assert contract["contractNo"] == "HT-2026-100"
    assert contract["document"]["version"] == 1
    assert contract["document"]["sha256"]
    assert list_contracts(database)[0]["title"] == "软件采购合同"
    downloaded = get_contract_document(database, contract["id"], 1)
    assert downloaded["fileName"] == "采购合同.docx"
    assert downloaded["content"] == _contract_docx()


def test_non_word_contract_is_rejected(tmp_path: Path) -> None:
    database = tmp_path / "contracts.sqlite3"
    initialize_platform_db(database)

    try:
        create_contract(database, "contract.pdf", b"pdf")
    except ValueError as error:
        assert "docx" in str(error)
    else:
        raise AssertionError("非 Word 合同必须被拒绝")


def test_new_word_version_updates_structured_facts_without_losing_original(tmp_path: Path) -> None:
    database = tmp_path / "contracts.sqlite3"
    initialize_platform_db(database)
    original = _contract_docx()
    contract = create_contract(database, "采购合同.docx", original)

    updated = add_contract_version(database, contract["id"], "采购合同-v2.docx", _contract_docx("HT-2026-100"), "reviewer")

    assert updated["currentVersion"] == 2
    assert updated["document"]["createdBy"] == "reviewer"
    assert get_contract_document(database, contract["id"], 1)["content"] == original


def test_comparison_exposes_rag_and_ontology_evidence(tmp_path: Path) -> None:
    database = tmp_path / "contracts.sqlite3"
    initialize_platform_db(database)
    contract = create_contract(database, "采购合同.docx", _contract_docx())

    result = compare_contract_engines(database, contract["id"], "这份合同的金额和甲乙方是什么？")

    assert result["rag"]["method"] == "lexical_retrieval"
    assert result["rag"]["citations"]
    assert result["ontology"]["method"] == "semantic_reasoning"
    assert result["ontology"]["facts"]["amount"] == 100000.0
    assert result["differences"][0]["dimension"] == "answer_basis"
