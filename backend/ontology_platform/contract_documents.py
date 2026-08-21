from __future__ import annotations

import hashlib
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any

# Chinese rule pattern for free-form paragraph extraction
RULE_IF_PATTERN = re.compile(
    r"(?:当|若|如果|如)\s*(?P<condition>[^，,。]+?)\s*(?:时|，|,)\s*(?:则|需|应|必须|不得|禁止|可以)?\s*(?P<action>[^。]+)",
)
RULE_MUST_PATTERN = re.compile(
    r"(?P<subject>[^，,。]{1,20})(?:需|应|必须|不得|禁止|可以)\s*(?P<action>[^。]+)"
)
RULE_SEVERITY_KEYWORDS = {
    "blocking": ["必须", "不得", "禁止", "严禁"],
    "warning": ["应", "应当", "不应", "不宜", "需", "需要", "建议", "推荐"],
    "info": ["可以", "可", "宜", "允许"],
}


def _convert_doc_to_docx(content: bytes) -> tuple[bytes, str]:
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_in:
        tmp_in.write(content)
        doc_path = Path(tmp_in.name)

    docx_path = doc_path.with_suffix(".docx")
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(doc_path.parent), str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode != 0 or not docx_path.exists():
            raise RuntimeError(f"LibreOffice 转换失败: {result.stderr.strip()}")
        return docx_path.read_bytes(), docx_path.name
    except FileNotFoundError as error:
        raise ValueError(
            "未检测到 LibreOffice，无法自动转换 .doc 文件。"
            "请手动将 .doc 另存为 .docx 格式，或安装 LibreOffice 后重试。"
        ) from error
    except subprocess.TimeoutExpired as error:
        raise ValueError(
            "LibreOffice 转换超时（30秒），请手动将 .doc 转换为 .docx 格式后重试。"
        ) from error
    finally:
        doc_path.unlink(missing_ok=True)
        docx_path.unlink(missing_ok=True)


def _ensure_docx(file_name: str, content: bytes) -> tuple[bytes, str]:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".docx":
        return content, file_name
    if suffix == ".doc":
        return _convert_doc_to_docx(content)
    raise ValueError(f"不支持的文件格式: {suffix}。当前支持 .docx 和 .doc。")


def parse_rule_docx_bytes(file_name: str, content: bytes, default_scope: str = "") -> dict[str, Any]:
    """Parse business rules from a Word document.

    `default_scope` is the business object assumed when a rule row omits its
    scope. The caller derives it from the target ontology, so no industry
    specific object code is baked into the parser.
    """
    content, safe_name = _ensure_docx(file_name, content)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)
    try:
        return parse_rule_docx(tmp_path, file_name=safe_name, content=content, default_scope=default_scope)
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_rule_docx(
    path: Path | str,
    file_name: str | None = None,
    content: bytes | None = None,
    default_scope: str = "",
) -> dict[str, Any]:
    try:
        from docx import Document
    except Exception as error:  # pragma: no cover
        raise ValueError("缺少 python-docx 依赖，无法解析 Word 规则文档。请安装 requirements.txt。") from error

    doc = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    tables = _extract_tables(doc)
    table_rules = _extract_rules_from_tables(tables, default_scope)
    paragraph_rules = _extract_rules_from_paragraphs(paragraphs, default_scope)
    # A structured document may naturally repeat words such as “必须” in rule
    # names and explanations. Do not reinterpret those fields as extra draft rules.
    free_text_rules = [] if table_rules or paragraph_rules else _extract_rules_from_free_text(paragraphs, default_scope)
    rules = _deduplicate_rules([*table_rules, *paragraph_rules, *free_text_rules])
    if not rules:
        raise ValueError("未从 Word 文档中识别到规则。建议使用表头：规则编码、规则名称、适用对象、规则类型、规则表达式、严重程度、自然语言说明。")
    file_bytes = content if content is not None else Path(path).read_bytes()
    return {
        "file": {
            "name": file_name or Path(path).name,
            "size": len(file_bytes),
            "md5": hashlib.md5(file_bytes).hexdigest(),
        },
        "rules": rules,
        "warnings": _rule_warnings(rules),
        "text": "\n".join(paragraphs),
    }


def _extract_tables(doc: Any) -> list[dict[str, Any]]:
    output = []
    for index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        output.append({"index": index, "rows": rows})
    return output


def _extract_rules_from_tables(tables: list[dict[str, Any]], default_scope: str = "") -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    for table in tables:
        if not table["rows"]:
            continue
        headers = [_normalize_header(value) for value in table["rows"][0]]
        if "code" not in headers or "expression" not in headers:
            continue
        for row in table["rows"][1:]:
            values = {headers[index]: row[index].strip() for index in range(min(len(headers), len(row))) if headers[index]}
            rule = _normalize_rule(values, default_scope)
            if rule:
                rules.append(rule)
    return rules


def _extract_rules_from_paragraphs(paragraphs: list[str], default_scope: str = "") -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    current: dict[str, str] = {}
    for paragraph in paragraphs:
        key, value = _split_rule_field(paragraph)
        if key == "code" and current.get("code"):
            rule = _normalize_rule(current, default_scope)
            if rule:
                rules.append(rule)
            current = {}
        if key:
            current[key] = value
    rule = _normalize_rule(current, default_scope)
    if rule:
        rules.append(rule)
    return rules


def _extract_rules_from_free_text(paragraphs: list[str], default_scope: str = "") -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    seen_codes: set[str] = set()

    for paragraph in paragraphs:
        for match in RULE_IF_PATTERN.finditer(paragraph):
            condition = match.group("condition").strip()
            action = match.group("action").strip()
            if not condition or not action:
                continue
            severity = _infer_severity(paragraph)
            scope = _infer_scope(default_scope, paragraph, condition, action)
            code = _slug_code(f"rule_{len(rules) + 1}")
            if code in seen_codes:
                code = _slug_code(f"rule_{len(rules) + 1}_{hash(condition) % 1000}")
            seen_codes.add(code)
            rules.append({
                "code": code,
                "name": f"规则_{len(rules) + 1}",
                "ruleType": "validation",
                "scopeObjectCode": scope,
                "expression": _condition_to_expression(condition),
                "severity": severity,
                "naturalLanguage": f"如果{condition}，则{action}",
                "status": "draft",
            })

        for match in RULE_MUST_PATTERN.finditer(paragraph):
            subject = match.group("subject").strip()
            action = match.group("action").strip()
            if not subject or not action:
                continue
            severity = _infer_severity(paragraph)
            scope = _infer_scope(paragraph, subject, action)
            code = _slug_code(f"must_{len(rules) + 1}")
            if code in seen_codes:
                code = _slug_code(f"must_{len(rules) + 1}_{hash(subject) % 1000}")
            seen_codes.add(code)
            rules.append({
                "code": code,
                "name": f"必须规则_{len(rules) + 1}",
                "ruleType": "validation",
                "scopeObjectCode": scope,
                "expression": _condition_to_expression(f"{subject}未{action}"),
                "severity": severity,
                "naturalLanguage": f"{subject}需{action}",
                "status": "draft",
            })

    return rules


def _infer_severity(text: str) -> str:
    for severity, keywords in RULE_SEVERITY_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            return severity
    return "warning"


def _infer_scope(default_scope: str, *hints: str) -> str:
    """Resolve the business object a free-text rule applies to.

    Free text carries no reliable scope, so the caller supplied default (taken
    from the target ontology) wins. Returning an empty string is preferred over
    guessing a business object that may not exist: the importer surfaces it as a
    warning instead of silently attaching the rule to the wrong object.
    """
    return _slug_code(default_scope)


def _condition_to_expression(condition: str) -> str:
    replacements = [
        ("大于", " > "), ("小于", " < "), ("等于", " == "),
        ("不为空", " != null"), ("为空", " == null"),
        ("不包含", " not in "), ("包含", " in "),
        ("超过", " > "), ("未超过", " <= "),
        ("达到", " >= "), ("未达到", " < "),
        ("属于", " in "), ("不属于", " not in "),
        ("是", " == "), ("不是", " != "),
        ("且", " and "), ("或", " or "), ("并", " and "),
    ]
    expr = condition
    for zh, op in replacements:
        expr = expr.replace(zh, op)
    return expr.strip() or condition


def _split_rule_field(text: str) -> tuple[str, str]:
    match = re.match(r"^([^:：]{2,20})[:：]\s*(.+)$", text)
    if not match:
        return "", ""
    return _normalize_header(match.group(1)), match.group(2).strip()


def _normalize_header(value: str) -> str:
    normalized = re.sub(r"\s+", "", value.strip().lower())
    aliases = {
        "规则编码": "code",
        "编码": "code",
        "code": "code",
        "规则名称": "name",
        "名称": "name",
        "name": "name",
        "规则类型": "rule_type",
        "类型": "rule_type",
        "ruletype": "rule_type",
        "rule_type": "rule_type",
        "适用对象": "scope_object_code",
        "业务对象": "scope_object_code",
        "对象": "scope_object_code",
        "scope": "scope_object_code",
        "scopeobjectcode": "scope_object_code",
        "规则表达式": "expression",
        "表达式": "expression",
        "expression": "expression",
        "严重程度": "severity",
        "级别": "severity",
        "severity": "severity",
        "自然语言说明": "natural_language",
        "业务说明": "natural_language",
        "说明": "natural_language",
        "naturallanguage": "natural_language",
        "状态": "status",
        "status": "status",
    }
    return aliases.get(normalized, "")


def _normalize_rule(values: dict[str, str], default_scope: str = "") -> dict[str, str] | None:
    code = _slug_code(values.get("code", ""))
    expression = values.get("expression", "").strip()
    # Fall back to the caller supplied scope (derived from the ontology) rather
    # than to any built-in business object.
    scope = _slug_code(values.get("scope_object_code", "")) or _slug_code(default_scope)
    if not code or not expression:
        return None
    return {
        "code": code,
        "name": values.get("name", "").strip() or code.replace("_", " ").title(),
        "ruleType": _normalize_rule_type(values.get("rule_type", "")),
        "scopeObjectCode": scope,
        "expression": expression,
        "severity": _normalize_severity(values.get("severity", "")),
        "naturalLanguage": values.get("natural_language", "").strip() or f"{code}：{expression}",
        "status": _normalize_status(values.get("status", "")),
    }


def _normalize_rule_type(value: str) -> str:
    normalized = value.strip().lower()
    mapping = {
        "校验": "validation",
        "验证": "validation",
        "validation": "validation",
        "派生": "derivation",
        "推导": "derivation",
        "derivation": "derivation",
        "状态流转": "transition",
        "流转": "transition",
        "transition": "transition",
        "风险": "risk",
        "risk": "risk",
        "建议": "recommendation",
        "recommendation": "recommendation",
        "权限": "permission",
        "permission": "permission",
    }
    return mapping.get(normalized, "validation")


def _normalize_severity(value: str) -> str:
    normalized = value.strip().lower()
    mapping = {
        "阻断": "blocking",
        "严重": "blocking",
        "blocking": "blocking",
        "警告": "warning",
        "warning": "warning",
        "提示": "info",
        "信息": "info",
        "info": "info",
    }
    return mapping.get(normalized, "warning")


def _normalize_status(value: str) -> str:
    normalized = value.strip().lower()
    mapping = {
        "已发布": "published",
        "发布": "published",
        "published": "published",
        "草稿": "draft",
        "draft": "draft",
        "停用": "disabled",
        "disabled": "disabled",
    }
    return mapping.get(normalized, "published")


def _slug_code(value: str) -> str:
    text = value.strip()
    if re.fullmatch(r"[A-Za-z][A-Za-z0-9_]*", text):
        return text
    slug = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_").lower()
    return slug


def _deduplicate_rules(rules: list[dict[str, str]]) -> list[dict[str, str]]:
    output: dict[str, dict[str, str]] = {}
    for rule in rules:
        output[rule["code"]] = rule
    return list(output.values())


def _rule_warnings(rules: list[dict[str, str]]) -> list[str]:
    warnings = []
    for rule in rules:
        if rule["status"] != "published":
            warnings.append(f"{rule['code']} 的状态不是 published，导入后可能不会参与已发布规则研判。")
        if not rule["scopeObjectCode"]:
            warnings.append(f"{rule['code']} 未指定适用业务对象，请在文档中补充“适用对象”列或在导入时指定默认对象。")
        if "." in rule["scopeObjectCode"]:
            warnings.append(f"{rule['code']} 的适用对象看起来像属性路径，请确认应为业务对象编码。")
    return warnings
