"""Document knowledge layer: text evidence anchored to the ontology.

This closes the third clause of the product's core claim:

    「订单 A123 不满足退款条件，因签收已超 15 天（规则 refund_window_check），
      依据《售后政策》第 3.2 条。」

"超 15 天" comes from a structured query and "不满足" from the rule engine; both
already worked. "第 3.2 条" is what this module supplies.

Design constraints from ADR-0009, each load-bearing:

- **Anchored, not free-floating.** A knowledge entry declares which business
  object or rule it serves. Retrieval narrows by that anchor *before* ranking by
  similarity, which is the opposite of "search everything and see what hits". An
  unanchored passage cannot answer "why does this text support this verdict", so
  it is not admissible evidence.
- **Governed, not auto-live.** Entries follow the existing pending/confirmed
  workflow. A mis-split clause that silently becomes judgement evidence produces
  a verdict that looks sourced but is not -- worse than no citation at all.
- **Citable.** Every entry carries a human-quotable locator ("第 3.2 条"), because
  an operator has to be able to look it up in the original document.

Chunking splits on clause boundaries rather than a fixed character window: a
citation is only meaningful if the chunk corresponds to something a person can
point at in the source.
"""

from __future__ import annotations

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from typing import Any, Iterable, Optional

from .context import PlatformDb
from .database import connect, last_insert_id
from .retrieval import tokenize
from .schema import SchemaBundle

logger = logging.getLogger(__name__)

# Clause markers seen in Chinese contracts, policies and regulations, ordered
# from most to least specific. These are *structural* patterns (numbering
# conventions), not business vocabulary, so they do not violate the
# domain-neutrality guarantee in ADR-0003.
CLAUSE_PATTERNS: tuple[tuple[str, str], ...] = (
    # 第3.2条 / 第 3.2 条
    (r"第\s*([0-9]+(?:\.[0-9]+)*)\s*条", "第{}条"),
    # 3.2 / 3.2.1 at the start of a line
    (r"^\s*([0-9]+(?:\.[0-9]+){1,})\s*[、.：:]?\s*", "{}"),
    # 一、二、 at the start of a line
    (r"^\s*([一二三四五六七八九十百]+)\s*[、.：:]", "第{}条"),
    # Article 5 / Section 5
    (r"^\s*(?:Article|Section)\s+([0-9]+(?:\.[0-9]+)*)", "Article {}"),
)

MIN_CHUNK_CHARS = 12
MAX_CHUNK_CHARS = 1200


@dataclass
class DocumentChunk:
    """One citable passage."""

    ordinal: int
    citation: str
    text: str
    heading: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "ordinal": self.ordinal,
            "citation": self.citation,
            "text": self.text,
            "heading": self.heading,
        }


@dataclass
class ParsedDocument:
    title: str
    chunks: list[DocumentChunk] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _match_clause(line: str) -> Optional[str]:
    """Return a citation label when a line opens a new clause."""
    for pattern, template in CLAUSE_PATTERNS:
        match = re.search(pattern, line) if pattern.startswith("第") else re.match(pattern, line)
        if match:
            return template.format(match.group(1))
    return None


def _looks_like_heading(line: str) -> bool:
    """A short line with no sentence-ending punctuation reads as a heading."""
    stripped = line.strip()
    if not stripped or len(stripped) > 40:
        return False
    return not re.search(r"[。；!?.;]$", stripped)


def split_into_chunks(text: str, *, title: str = "") -> ParsedDocument:
    """Split document text into citable chunks on clause boundaries.

    Falls back to paragraph grouping when a document carries no recognisable
    numbering; in that case citations are positional ("第 3 段"), which is weaker
    but still points a reader somewhere specific.
    """
    document = ParsedDocument(title=title)
    lines = [line.strip() for line in (text or "").replace("\r\n", "\n").split("\n")]
    lines = [line for line in lines if line]
    if not lines:
        document.warnings.append("文档没有可提取的文本内容。")
        return document

    current_citation = ""
    current_heading = ""
    buffer: list[str] = []
    ordinal = 0
    numbered = False

    def flush() -> None:
        nonlocal buffer, ordinal
        body = "\n".join(buffer).strip()
        buffer = []
        if len(body) < MIN_CHUNK_CHARS:
            return
        ordinal += 1
        # Over-long clauses are truncated rather than silently split, so the
        # citation keeps pointing at exactly one clause.
        if len(body) > MAX_CHUNK_CHARS:
            document.warnings.append(f"{current_citation or f'第 {ordinal} 段'} 超过 {MAX_CHUNK_CHARS} 字，已截断。")
            body = body[:MAX_CHUNK_CHARS]
        document.chunks.append(
            DocumentChunk(
                ordinal=ordinal,
                citation=current_citation or f"第 {ordinal} 段",
                text=body,
                heading=current_heading,
            )
        )

    for line in lines:
        citation = _match_clause(line)
        if citation:
            numbered = True
            flush()
            current_citation = citation
            buffer.append(line)
            continue
        if _looks_like_heading(line) and not buffer:
            current_heading = line
            continue
        buffer.append(line)
    flush()

    if not numbered:
        document.warnings.append("未识别到条款编号，已按段落切分，引用定位为段序号而非条款号。")
    if not document.chunks:
        document.warnings.append("切分后没有达到最小长度的条目。")
    return document


def extract_text_from_docx(content: bytes) -> str:
    """Extract paragraph text from a .docx payload."""
    try:
        import io

        from docx import Document as DocxDocument
    except ImportError as error:  # pragma: no cover - dependency is declared
        raise ValueError("解析 Word 文档需要安装 python-docx。") from error
    document = DocxDocument(io.BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def init_knowledge_schema(conn: Any) -> None:
    """Create the document and chunk tables.

    Declared here rather than in database.SCHEMA_DEFINITIONS to keep the feature
    self-contained; called from application startup alongside the other optional
    schemas.
    """
    SchemaBundle(name="knowledge_documents", tables=KNOWLEDGE_SCHEMA).apply(conn)


KNOWLEDGE_SCHEMA: tuple[dict[str, str], ...] = (
    {
        "sqlite": """
        create table if not exists knowledge_document (
            id integer primary key autoincrement,
            ontology_id integer not null references ontology(id),
            title text not null,
            source_name text not null default '',
            content_hash text not null default '',
            chunk_count integer not null default 0,
            status text not null default 'pending',
            uploaded_by text not null default '',
            created_at text not null default current_timestamp,
            unique(ontology_id, content_hash)
        )""",
        "postgresql": """
        create table if not exists knowledge_document (
            id serial primary key,
            ontology_id integer not null references ontology(id),
            title text not null,
            source_name text not null default '',
            content_hash text not null default '',
            chunk_count integer not null default 0,
            status text not null default 'pending',
            uploaded_by text not null default '',
            created_at timestamp not null default current_timestamp,
            unique(ontology_id, content_hash)
        )""",
        "mysql": """
        create table if not exists knowledge_document (
            id integer primary key auto_increment,
            ontology_id integer not null,
            title varchar(500) not null,
            source_name varchar(500) not null default '',
            content_hash varchar(64) not null default '',
            chunk_count integer not null default 0,
            status varchar(50) not null default 'pending',
            uploaded_by varchar(255) not null default '',
            created_at datetime not null default current_timestamp,
            unique key uniq_doc_hash (ontology_id, content_hash)
        )""",
    },
    {
        "sqlite": """
        create table if not exists knowledge_entry (
            id integer primary key autoincrement,
            document_id integer not null references knowledge_document(id),
            ontology_id integer not null references ontology(id),
            ordinal integer not null,
            citation text not null,
            heading text not null default '',
            content text not null,
            object_code text not null default '',
            rule_code text not null default '',
            status text not null default 'pending',
            reviewer text not null default '',
            reviewed_at text,
            token_summary text not null default '{}',
            created_at text not null default current_timestamp
        )""",
        "postgresql": """
        create table if not exists knowledge_entry (
            id serial primary key,
            document_id integer not null references knowledge_document(id),
            ontology_id integer not null references ontology(id),
            ordinal integer not null,
            citation text not null,
            heading text not null default '',
            content text not null,
            object_code text not null default '',
            rule_code text not null default '',
            status text not null default 'pending',
            reviewer text not null default '',
            reviewed_at timestamp,
            token_summary text not null default '{}',
            created_at timestamp not null default current_timestamp
        )""",
        "mysql": """
        create table if not exists knowledge_entry (
            id integer primary key auto_increment,
            document_id integer not null,
            ontology_id integer not null,
            ordinal integer not null,
            citation varchar(255) not null,
            heading varchar(500) not null default '',
            content text not null,
            object_code varchar(255) not null default '',
            rule_code varchar(255) not null default '',
            status varchar(50) not null default 'pending',
            reviewer varchar(255) not null default '',
            reviewed_at datetime,
            token_summary text,
            created_at datetime not null default current_timestamp
        )""",
    },
)


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def ingest_document(
    platform_db: PlatformDb,
    ontology_id: int,
    title: str,
    text: str,
    *,
    source_name: str = "",
    object_code: str = "",
    rule_code: str = "",
    actor: str = "system",
) -> dict[str, Any]:
    """Split a document into pending knowledge entries.

    Entries land as `pending`: nothing becomes judgement evidence until a human
    confirms it, and the anchor (object_code / rule_code) can be corrected during
    review.
    """
    if not (title or "").strip():
        raise ValueError("文档标题不能为空")
    parsed = split_into_chunks(text, title=title)
    if not parsed.chunks:
        raise ValueError("文档切分后没有可用条目：" + "；".join(parsed.warnings))

    digest = _content_hash(text)
    with connect(platform_db) as conn:
        ontology = conn.execute("select id, status from ontology where id = ?", (ontology_id,)).fetchone()
        if ontology is None:
            raise ValueError(f"本体不存在: {ontology_id}")
        if ontology["status"] == "published":
            raise ValueError("已发布本体不可新增知识条目，请派生新版本。")

        existing = conn.execute(
            "select id from knowledge_document where ontology_id = ? and content_hash = ?",
            (ontology_id, digest),
        ).fetchone()
        if existing is not None:
            raise ValueError(f"该文档内容已导入（文档 #{existing['id']}），请勿重复上传。")

        conn.execute(
            """
            insert into knowledge_document
                (ontology_id, title, source_name, content_hash, chunk_count, status, uploaded_by)
            values (?, ?, ?, ?, ?, 'pending', ?)
            """,
            (ontology_id, title, source_name, digest, len(parsed.chunks), actor),
        )
        document_id = last_insert_id(conn)

        for chunk in parsed.chunks:
            conn.execute(
                """
                insert into knowledge_entry
                    (document_id, ontology_id, ordinal, citation, heading, content,
                     object_code, rule_code, status, token_summary)
                values (?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?)
                """,
                (
                    document_id,
                    ontology_id,
                    chunk.ordinal,
                    chunk.citation,
                    chunk.heading,
                    chunk.text,
                    object_code,
                    rule_code,
                    json.dumps(summarize_tokens(chunk.text), ensure_ascii=False),
                ),
            )

        conn.execute(
            "insert into audit_log (actor, action, target_type, target_id, detail) values (?, ?, ?, ?, ?)",
            (
                actor,
                "ingest_knowledge_document",
                "knowledge_document",
                str(document_id),
                json.dumps(
                    {"title": title, "chunks": len(parsed.chunks), "objectCode": object_code},
                    ensure_ascii=False,
                ),
            ),
        )

    return {
        "documentId": document_id,
        "title": title,
        "chunkCount": len(parsed.chunks),
        "status": "pending",
        "warnings": parsed.warnings,
        "citations": [chunk.citation for chunk in parsed.chunks],
        "note": "条目为 pending 状态，确认后才会作为判定依据被检索。",
    }


def summarize_tokens(text: str) -> dict[str, int]:
    """Term frequencies used by the default retrieval backend.

    Stored at ingest time so retrieval does not re-tokenise every entry on every
    query. See retrieval.py for why the default tokenizer works this way.
    """
    counts: dict[str, int] = {}
    for token in tokenize(text):
        counts[token] = counts.get(token, 0) + 1
    return counts


def review_knowledge_entry(
    platform_db: PlatformDb,
    entry_id: int,
    status: str,
    *,
    object_code: Optional[str] = None,
    rule_code: Optional[str] = None,
    reviewer: str = "system",
) -> dict[str, Any]:
    """Confirm or reject an entry, optionally correcting its anchor."""
    if status not in {"confirmed", "rejected", "pending"}:
        raise ValueError(f"不支持的状态: {status}")
    with connect(platform_db) as conn:
        entry = conn.execute("select * from knowledge_entry where id = ?", (entry_id,)).fetchone()
        if entry is None:
            raise ValueError(f"知识条目不存在: {entry_id}")

        resolved_object = entry["object_code"] if object_code is None else object_code
        resolved_rule = entry["rule_code"] if rule_code is None else rule_code

        # A confirmed entry with no anchor would be free-floating text, which is
        # exactly what ADR-0009 rules out.
        if status == "confirmed" and not (resolved_object or resolved_rule):
            raise ValueError("确认知识条目前必须指定锚定的业务对象或规则。")

        conn.execute(
            """
            update knowledge_entry
            set status = ?, object_code = ?, rule_code = ?, reviewer = ?, reviewed_at = current_timestamp
            where id = ?
            """,
            (status, resolved_object, resolved_rule, reviewer, entry_id),
        )
        conn.execute(
            "insert into audit_log (actor, action, target_type, target_id, detail) values (?, ?, ?, ?, ?)",
            (
                reviewer,
                "review_knowledge_entry",
                "knowledge_entry",
                str(entry_id),
                json.dumps(
                    {"status": status, "objectCode": resolved_object, "ruleCode": resolved_rule},
                    ensure_ascii=False,
                ),
            ),
        )
        _refresh_document_status(conn, int(entry["document_id"]))
    return {
        "id": entry_id,
        "status": status,
        "objectCode": resolved_object,
        "ruleCode": resolved_rule,
    }


def _refresh_document_status(conn: Any, document_id: int) -> None:
    rows = conn.execute(
        "select status, count(*) as count from knowledge_entry where document_id = ? group by status",
        (document_id,),
    ).fetchall()
    counts = {row["status"]: int(row["count"]) for row in rows}
    pending = counts.get("pending", 0)
    confirmed = counts.get("confirmed", 0)
    status = "pending" if pending else ("confirmed" if confirmed else "rejected")
    conn.execute("update knowledge_document set status = ? where id = ?", (status, document_id))


def list_documents(platform_db: PlatformDb, ontology_id: int) -> list[dict[str, Any]]:
    with connect(platform_db) as conn:
        rows = conn.execute(
            """
            select d.id, d.title, d.source_name, d.chunk_count, d.status, d.uploaded_by, d.created_at,
                   (select count(*) from knowledge_entry e
                    where e.document_id = d.id and e.status = 'confirmed') as confirmed_count,
                   (select count(*) from knowledge_entry e
                    where e.document_id = d.id and e.status = 'pending') as pending_count
            from knowledge_document d
            where d.ontology_id = ?
            order by d.id desc
            """,
            (ontology_id,),
        ).fetchall()
    return [
        {
            "id": row["id"],
            "title": row["title"],
            "sourceName": row["source_name"],
            "chunkCount": int(row["chunk_count"]),
            "confirmedCount": int(row["confirmed_count"] or 0),
            "pendingCount": int(row["pending_count"] or 0),
            "status": row["status"],
            "uploadedBy": row["uploaded_by"],
            "createdAt": str(row["created_at"]),
        }
        for row in rows
    ]


def list_entries(
    platform_db: PlatformDb,
    ontology_id: int,
    *,
    document_id: Optional[int] = None,
    status: Optional[str] = None,
    limit: int = 100,
) -> list[dict[str, Any]]:
    clauses = ["ontology_id = ?"]
    params: list[Any] = [ontology_id]
    if document_id is not None:
        clauses.append("document_id = ?")
        params.append(document_id)
    if status:
        clauses.append("status = ?")
        params.append(status)
    params.append(max(1, min(int(limit), 500)))
    with connect(platform_db) as conn:
        rows = conn.execute(
            f"""
            select id, document_id, ordinal, citation, heading, content,
                   object_code, rule_code, status, reviewer
            from knowledge_entry
            where {" and ".join(clauses)}
            order by document_id desc, ordinal
            limit ?
            """,
            tuple(params),
        ).fetchall()
    return [_entry_dict(row) for row in rows]


def _entry_dict(row: Any) -> dict[str, Any]:
    return {
        "id": row["id"],
        "documentId": row["document_id"],
        "ordinal": int(row["ordinal"]),
        "citation": row["citation"],
        "heading": row["heading"],
        "content": row["content"],
        "objectCode": row["object_code"],
        "ruleCode": row["rule_code"],
        "status": row["status"],
        "reviewer": row["reviewer"],
    }


def load_confirmed_entries(
    conn: Any,
    ontology_id: int,
    *,
    object_codes: Iterable[str] = (),
    rule_codes: Iterable[str] = (),
) -> list[dict[str, Any]]:
    """Confirmed entries anchored to the given objects or rules.

    Narrowing by anchor happens here, in SQL, before any similarity ranking --
    that ordering is the whole point of ADR-0009. An empty anchor set means the
    caller wants every confirmed entry for the ontology.
    """
    objects = [code for code in object_codes if code]
    rules = [code for code in rule_codes if code]
    # Qualify every column: this query joins knowledge_document, which also has
    # an ontology_id.
    clauses = ["e.ontology_id = ?", "e.status = 'confirmed'"]
    params: list[Any] = [ontology_id]
    if objects or rules:
        anchor_clauses = []
        if objects:
            anchor_clauses.append(f"e.object_code in ({', '.join('?' for _ in objects)})")
            params.extend(objects)
        if rules:
            anchor_clauses.append(f"e.rule_code in ({', '.join('?' for _ in rules)})")
            params.extend(rules)
        clauses.append(f"({' or '.join(anchor_clauses)})")
    rows = conn.execute(
        f"""
        select e.id, e.document_id, e.ordinal, e.citation, e.heading, e.content,
               e.object_code, e.rule_code, e.status, e.reviewer, e.token_summary,
               d.title as document_title
        from knowledge_entry e
        join knowledge_document d on d.id = e.document_id
        where {" and ".join(clauses)}
        order by e.document_id, e.ordinal
        """,
        tuple(params),
    ).fetchall()
    entries = []
    for row in rows:
        entry = _entry_dict(row)
        entry["documentTitle"] = row["document_title"]
        try:
            entry["tokenSummary"] = json.loads(row["token_summary"] or "{}")
        except (TypeError, ValueError):
            entry["tokenSummary"] = {}
        entries.append(entry)
    return entries
